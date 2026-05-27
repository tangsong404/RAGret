from __future__ import annotations

import io
import types
from pathlib import Path

import pytest

from ragret.preprocess.docx import preprocess_docx
from ragret.preprocess.pdf import preprocess_pdf


def test_preprocess_pdf_writes_back_image_blocks(monkeypatch, tmp_path: Path) -> None:
    class _Img:
        name = "img-1"
        data = b"\x89PNG\r\n\x1a\nxx"

    class _Page:
        images = [_Img()]

        @staticmethod
        def extract_text() -> str:
            return "hello page"

    class _Reader:
        pages = [_Page()]

        def __init__(self, _path: str) -> None:
            pass

    monkeypatch.setitem(__import__("sys").modules, "pypdf", types.SimpleNamespace(PdfReader=_Reader))
    p = tmp_path / "a.pdf"
    p.write_bytes(b"%PDF-1.4")
    out = preprocess_pdf(p, image_handler=lambda _b, _n: "流程图\n\n图片: /api/kb/x/assets/y.png")
    assert "hello page" in out
    assert "流程图" in out
    assert "图片: /api/kb/x/assets/y.png" in out


def test_preprocess_docx_inline_image_order(tmp_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Inches
        from PIL import Image
    except ImportError:
        pytest.skip("python-docx and Pillow required")

    buf = io.BytesIO()
    Image.new("RGB", (12, 12), color="blue").save(buf, format="PNG")
    buf.seek(0)

    doc = Document()
    doc.add_paragraph("段落一")
    mid = doc.add_paragraph()
    mid.add_run("中间文字")
    mid.add_run().add_picture(buf, width=Inches(0.5))
    doc.add_paragraph("段落三")
    path = tmp_path / "inline.docx"
    doc.save(path)

    out = preprocess_docx(
        path,
        image_handler=lambda _b, _n, _c=None: "示意图\n\n图片: /api/kb/x/assets/i.png",
    )
    assert "--- embedded images ---" not in out
    i1 = out.index("段落一")
    i_img = out.index("示意图")
    i3 = out.index("段落三")
    assert i1 < i_img < i3
    assert "中间文字" in out
